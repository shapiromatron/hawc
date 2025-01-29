import json
import logging

import pandas as pd
from django.contrib.contenttypes.fields import GenericRelation
from django.core.exceptions import ValidationError
from django.db import models
from django.urls import reverse
from django.utils import timezone
from docx import Document as create_document
from plotly.graph_objs._figure import Figure
from plotly.io import from_json
from pydantic import BaseModel as PydanticModel
from reversion import revisions as reversion

from ...tools.tables.ept import EvidenceProfileTable
from ...tools.tables.generic import BaseTable, GenericTable
from ...tools.tables.parser import QuillParser
from ...tools.tables.set import StudyEvaluationTable
from ..animal.exports import EndpointFlatDataPivot, EndpointGroupFlatDataPivot
from ..animal.models import Endpoint
from ..assessment.constants import EpiVersion, RobName
from ..assessment.models import Assessment, BaseEndpoint, DoseUnits, LabeledItem
from ..common.helper import FlatExport, PydanticToDjangoError, ReportExport, SerializerHelper
from ..common.validators import validate_html_tags, validate_hyperlinks
from ..eco.exports import EcoFlatComplete
from ..epi.exports import OutcomeDataPivot
from ..epimeta.exports import MetaResultFlatDataPivot
from ..epiv2.exports import EpiFlatComplete
from ..invitro import exports as ivexports
from ..lit.models import Reference, ReferenceFilterTag, Search
from ..riskofbias.models import RiskOfBiasScore
from ..riskofbias.serializers import AssessmentRiskOfBiasSerializer
from ..study.models import Study
from . import constants, managers, prefilters
from .actions.rob import get_rob_visual_form_data

logger = logging.getLogger(__name__)


class SummaryTable(models.Model):
    objects = managers.SummaryTableManager()

    TABLE_SCHEMA_MAP = {
        constants.TableType.GENERIC: GenericTable,
        constants.TableType.EVIDENCE_PROFILE: EvidenceProfileTable,
        constants.TableType.STUDY_EVALUATION: StudyEvaluationTable,
    }

    assessment = models.ForeignKey(Assessment, on_delete=models.CASCADE)
    title = models.CharField(max_length=128)
    slug = models.SlugField(
        verbose_name="URL Name",
        help_text="The URL (web address) used to describe this object "
        "(no spaces or special-characters).",
    )
    content = models.JSONField(default=dict)
    table_type = models.PositiveSmallIntegerField(
        choices=constants.TableType, default=constants.TableType.GENERIC
    )
    published = models.BooleanField(
        default=False,
        verbose_name="Publish table for public viewing",
        help_text="For assessments marked for public viewing, mark table to be viewable by public",
    )
    caption = models.TextField(blank=True, validators=[validate_html_tags, validate_hyperlinks])
    labels = GenericRelation(LabeledItem, related_query_name="summary_tables")
    created = models.DateTimeField(auto_now_add=True)
    last_updated = models.DateTimeField(auto_now=True)

    BREADCRUMB_PARENT = "assessment"

    class Meta:
        unique_together = (
            ("assessment", "title"),
            ("assessment", "slug"),
        )

    def __str__(self):
        return self.title

    def get_assessment(self):
        return self.assessment

    @classmethod
    def get_list_url(cls, assessment_id: int):
        return reverse("summary:tables_list", args=(assessment_id,))

    @classmethod
    def get_api_list_url(cls, assessment_id: int):
        return reverse("summary:api:summary-table-list") + f"?assessment_id={assessment_id}"

    def get_absolute_url(self):
        return reverse(
            "summary:tables_detail",
            args=(
                self.assessment_id,
                self.slug,
            ),
        )

    def get_update_url(self):
        return reverse(
            "summary:tables_update",
            args=(
                self.assessment_id,
                self.slug,
            ),
        )

    def get_delete_url(self):
        return reverse(
            "summary:tables_delete",
            args=(
                self.assessment_id,
                self.slug,
            ),
        )

    def get_api_url(self):
        return reverse("summary:api:summary-table-detail", args=(self.id,))

    def get_api_word_url(self):
        return reverse("summary:api:summary-table-docx", args=(self.id,))

    def get_content_schema_class(self):
        if self.table_type not in self.TABLE_SCHEMA_MAP:
            raise ValueError(f"Table type not found: {self.table_type}")
        return self.TABLE_SCHEMA_MAP[self.table_type]

    def get_table(self) -> BaseTable:
        schema_class = self.get_content_schema_class()
        # ensure the assessment id is from the object; not custom config
        kwargs = {}
        if "assessment_id" in schema_class.model_json_schema()["properties"]:
            kwargs["assessment_id"] = self.assessment_id
        return schema_class.model_validate(dict(self.content, **kwargs))

    @classmethod
    def build_default(cls, assessment_id: int, table_type: int) -> "SummaryTable":
        """Build an incomplete, but default SummaryTable instance"""
        instance = cls(assessment_id=assessment_id, table_type=table_type)
        schema = instance.get_content_schema_class()
        instance.content = schema.get_default_props()
        return instance

    def to_docx(self, base_url: str = "", landscape: bool = True):
        docx = create_document()
        parser = QuillParser(base_url=base_url)
        docx.add_heading(self.title)
        url = self.get_absolute_url()
        parser.feed(f'<p><a href="{url}">View Online</a></p>', docx)
        table = self.get_table()
        table.to_docx(parser=parser, docx=docx, landscape=landscape)
        if self.caption:
            parser.feed(self.caption, docx)
        return ReportExport(docx=docx, filename=self.slug)

    @classmethod
    def get_data(cls, table_type: int, assessment_id: int, **kwargs):
        return cls.TABLE_SCHEMA_MAP[table_type].get_data(assessment_id=assessment_id, **kwargs)

    def clean(self):
        # make sure table can be built
        with PydanticToDjangoError(field="content"):
            try:
                self.get_table()
            except ValueError as e:
                raise ValidationError({"content": str(e)}) from e

        # clean up control characters before string validation
        content_str = json.dumps(self.content).replace('\\"', '"')

        # validate tags used in text
        validate_html_tags(content_str, "content")
        validate_hyperlinks(content_str, "content")


class HeatmapDataset(PydanticModel):
    type: str
    name: str
    url: str


class HeatmapDatasets(PydanticModel):
    datasets: list[HeatmapDataset]


class Visual(models.Model):
    objects = managers.VisualManager()

    title = models.CharField(max_length=128)
    slug = models.SlugField(
        verbose_name="URL Name",
        help_text="The URL (web address) used to describe this object "
        "(no spaces or special-characters).",
    )
    assessment = models.ForeignKey(Assessment, on_delete=models.CASCADE, related_name="visuals")
    visual_type = models.PositiveSmallIntegerField(choices=constants.VisualType)
    evidence_type = models.PositiveSmallIntegerField(choices=constants.StudyType)
    dose_units = models.ForeignKey(DoseUnits, on_delete=models.SET_NULL, blank=True, null=True)
    endpoints = models.ManyToManyField(
        BaseEndpoint,
        related_name="visuals",
        help_text="Endpoints to be included in visualization",
        blank=True,
    )
    settings = models.JSONField(default=dict)
    caption = models.TextField(blank=True)
    published = models.BooleanField(
        default=False,
        verbose_name="Publish visual for public viewing",
        help_text="For assessments marked for public viewing, mark visual to be viewable by public",
    )
    prefilters = models.JSONField(default=dict)
    image = models.ImageField(
        upload_to="summary/visual/images",
        blank=True,
        null=True,
        help_text="Upload an image file. Valid formats: png, jpg, jpeg. Must be > 10KB and < 3MB in size.",
    )
    dp_id = models.BigIntegerField(
        editable=False, blank=True, null=True, help_text="data pivot migration"
    )
    dp_slug = models.SlugField(editable=False, blank=True, help_text="data pivot migration")
    dataset = models.ForeignKey(
        "assessment.Dataset", on_delete=models.CASCADE, blank=True, null=True
    )
    created = models.DateTimeField(auto_now_add=True)
    last_updated = models.DateTimeField(auto_now=True)
    labels = GenericRelation(LabeledItem, related_query_name="visuals")

    BREADCRUMB_PARENT = "assessment"

    class Meta:
        unique_together = (("assessment", "slug"),)

    def __str__(self):
        return self.title

    @staticmethod
    def get_list_url(assessment_id):
        return reverse("summary:visualization_list", args=(assessment_id,))

    def get_absolute_url(self):
        return reverse("summary:visualization_detail", args=(self.assessment_id, self.slug))

    def get_update_url(self):
        return reverse("summary:visualization_update", args=(self.assessment_id, self.slug))

    def get_delete_url(self):
        return reverse("summary:visualization_delete", args=(self.assessment_id, self.slug))

    def get_assessment(self):
        return self.assessment

    def get_api_detail(self):
        return reverse("summary:api:visual-detail", args=(self.id,))

    def get_data_url(self):
        return reverse("summary:api:visual-data", args=(self.id,))

    def get_dp_update_settings(self):
        return reverse(
            "summary:visualization_update_settings", args=(self.assessment_id, self.slug)
        )

    def get_api_heatmap_datasets(self):
        return reverse("summary:api:assessment-heatmap-datasets", args=(self.assessment_id,))

    def get_visual_type_display(self):
        label = constants.VisualType(self.visual_type).label
        if (
            self.visual_type == constants.VisualType.ROB_HEATMAP
            or self.visual_type == constants.VisualType.ROB_BARCHART
        ) and self.assessment.rob_name == RobName.SE:
            label = label.replace("risk of bias", "study evaluation")
        elif self.visual_type == constants.VisualType.DATA_PIVOT_QUERY:
            if self.evidence_type == constants.StudyType.BIOASSAY:
                label = "Data pivot (animal bioassay)"
            elif self.evidence_type == constants.StudyType.EPI:
                label = "Data pivot (epidemiology)"
            elif self.evidence_type == constants.StudyType.EPI_META:
                label = "Data pivot (epidemiology meta-analysis/pooled-analysis)"
            elif self.evidence_type == constants.StudyType.IN_VITRO:
                label = "Data pivot (in vitro)"
            elif self.evidence_type == constants.StudyType.ECO:
                label = "Data pivot (ecology)"
        return label

    @property
    def is_data_pivot(self) -> bool:
        return self.visual_type in (
            constants.VisualType.DATA_PIVOT_QUERY,
            constants.VisualType.DATA_PIVOT_FILE,
        )

    @classmethod
    def get_heatmap_datasets(cls, assessment: Assessment) -> HeatmapDatasets:
        datasets = [
            HeatmapDataset(
                type="Literature",
                name="Literature summary",
                url=reverse("lit:api:assessment-tag-heatmap", args=(assessment.id,)),
            )
        ]
        if assessment.has_animal_data:
            datasets.extend(
                [
                    HeatmapDataset(
                        type="Bioassay",
                        name="Bioassay study design",
                        url=reverse("animal:api:assessment-study-heatmap", args=(assessment.id,)),
                    ),
                    HeatmapDataset(
                        type="Bioassay",
                        name="Bioassay study design (including unpublished HAWC data)",
                        url=reverse("animal:api:assessment-study-heatmap", args=(assessment.id,))
                        + "?unpublished=true",
                    ),
                    HeatmapDataset(
                        type="Bioassay",
                        name="Bioassay endpoint summary",
                        url=reverse(
                            "animal:api:assessment-endpoint-heatmap", args=(assessment.id,)
                        ),
                    ),
                    HeatmapDataset(
                        type="Bioassay",
                        name="Bioassay endpoint summary (including unpublished HAWC data)",
                        url=reverse("animal:api:assessment-endpoint-heatmap", args=(assessment.id,))
                        + "?unpublished=true",
                    ),
                    HeatmapDataset(
                        type="Bioassay",
                        name="Bioassay endpoint with doses",
                        url=reverse(
                            "animal:api:assessment-endpoint-doses-heatmap", args=(assessment.id,)
                        ),
                    ),
                    HeatmapDataset(
                        type="Bioassay",
                        name="Bioassay endpoint with doses (including unpublished HAWC data)",
                        url=reverse(
                            "animal:api:assessment-endpoint-doses-heatmap", args=(assessment.id,)
                        )
                        + "?unpublished=true",
                    ),
                ]
            )
        if assessment.has_epi_data:
            if assessment.epi_version == EpiVersion.V1:
                datasets.extend(
                    [
                        HeatmapDataset(
                            type="Epi",
                            name="Epidemiology study design",
                            url=reverse("epi:api:assessment-study-heatmap", args=(assessment.id,)),
                        ),
                        HeatmapDataset(
                            type="Epi",
                            name="Epidemiology study design (including unpublished HAWC data)",
                            url=reverse("epi:api:assessment-study-heatmap", args=(assessment.id,))
                            + "?unpublished=true",
                        ),
                        HeatmapDataset(
                            type="Epi",
                            name="Epidemiology result summary",
                            url=reverse("epi:api:assessment-result-heatmap", args=(assessment.id,)),
                        ),
                        HeatmapDataset(
                            type="Epi",
                            name="Epidemiology result summary (including unpublished HAWC data)",
                            url=reverse("epi:api:assessment-result-heatmap", args=(assessment.id,))
                            + "?unpublished=true",
                        ),
                    ]
                )
            elif assessment.epi_version == EpiVersion.V2:
                datasets.extend(
                    [
                        HeatmapDataset(
                            type="Epi",
                            name="Epidemiology evidence map",
                            url=reverse("epiv2:api:assessment-study-export", args=(assessment.id,)),
                        ),
                        HeatmapDataset(
                            type="Epi",
                            name="Epidemiology evidence map (including unpublished HAWC data)",
                            url=reverse("epiv2:api:assessment-study-export", args=(assessment.id,))
                            + "?unpublished=true",
                        ),
                        HeatmapDataset(
                            type="Epi",
                            name="Epidemiology data extractions",
                            url=reverse("epiv2:api:assessment-export", args=(assessment.id,)),
                        ),
                        HeatmapDataset(
                            type="Epi",
                            name="Epidemiology data extractions (including unpublished HAWC data)",
                            url=reverse("epiv2:api:assessment-export", args=(assessment.id,))
                            + "?unpublished=true",
                        ),
                    ]
                )
            else:
                raise ValueError("Unknown epi data type")
        additional_datasets = [
            HeatmapDataset(type="Dataset", name=f"Dataset: {ds.name}", url=ds.get_api_data_url())
            for ds in assessment.datasets.all()
        ]
        datasets.extend(additional_datasets)
        return HeatmapDatasets(datasets=datasets)

    def get_prisma_data(self) -> dict:
        return {
            "reference_tag_pairs": list(
                Reference.objects.tag_pairs(self.assessment.references.all().order_by("id"))
            ),
            "reference_search_pairs": list(
                Reference.searches.through.objects.filter(
                    reference_id__in=self.assessment.references.all()
                )
                .values("search_id", "reference_id")
                .order_by("search_id", "reference_id")
            ),
            "searches": list(
                Search.objects.filter(assessment_id=self.assessment.id)
                .values("id", "title")
                .order_by("id")
            ),
            "tags": list(
                ReferenceFilterTag.as_dataframe(self.assessment.id).to_dict(orient="records")
            ),
            "references": list(
                Reference.objects.filter(assessment=self.assessment.id)
                .values_list("id", flat=True)
                .order_by("id")
            ),
            "reference_detail_url": reverse("lit:interactive", args=(self.assessment.id,))
            + "?action=venn_reference_list",
        }

    def get_rob_data(self) -> dict:
        return get_rob_visual_form_data(
            assessment_id=self.assessment_id,
            study_type=constants.StudyType(self.evidence_type),
        )

    def get_data(self) -> dict:
        """Get data needed to display Visual."""
        match self.visual_type:
            case constants.VisualType.PRISMA:
                return self.get_prisma_data()
            case _:
                return {}

    def read_config(self) -> dict:
        """Configuration required to render an instance of the visual in read-only views."""
        match self.visual_type:
            case constants.VisualType.PRISMA:
                return dict(
                    settings=self.settings,
                    api_data_url=reverse("summary:api:visual-json-data", args=(self.id,)),
                )
            case _:
                return {}

    def get_filterset_class(self):
        return prefilters.get_prefilter_cls(self.visual_type, self.evidence_type, self.assessment)

    def get_filterset(self, data, assessment, **kwargs):
        return self.get_filterset_class()(data=data, assessment=assessment, **kwargs)

    def get_endpoints(self) -> models.QuerySet[Endpoint]:
        if self.visual_type == constants.VisualType.BIOASSAY_AGGREGATION:
            return Endpoint.objects.assessment_qs(self.assessment).filter(
                id__in=self.endpoints.values_list("id", flat=True)
            )
        elif self.visual_type == constants.VisualType.BIOASSAY_CROSSVIEW:
            filters = {"animal_group__dosing_regime__doses__dose_units_id": self.dose_units_id}
            fs = self.get_filterset(self.prefilters, self.assessment)
            form = fs.form
            fs.set_passthrough_options(form)
            fs.form.is_valid()
            return fs.qs.filter(**filters).distinct("id")

        return Endpoint.objects.none()

    def get_studies(self) -> models.QuerySet[Study]:
        """
        If there are endpoint-level prefilters, we get all studies which
        match this criteria. Otherwise, we use the M2M list of studies attached
        to the model.
        """
        qs = Study.objects.none()

        if self.visual_type in [
            constants.VisualType.ROB_HEATMAP,
            constants.VisualType.ROB_BARCHART,
        ]:
            fs = self.get_filterset(self.prefilters, self.assessment)
            form = fs.form
            fs.set_passthrough_options(form)
            fs.form.is_valid()
            qs = fs.qs

        return qs

    def get_editing_dataset(self, request):
        # Generate a pseudo-return when editing or creating a dataset.
        # Do not include the settings field; this will be set from the
        # input-form. Should approximately mirror the Visual API from rest-framework.

        dose_units = None
        try:
            dose_units = int(request.POST.get("dose_units"))
        except (TypeError, ValueError):
            # TypeError if dose_units is None; ValueError if dose_units is ""
            pass

        return {
            "assessment": self.assessment_id,
            "assessment_rob_name": self.assessment.get_rob_name_display(),
            "title": request.POST.get("title"),
            "slug": request.POST.get("slug"),
            "caption": request.POST.get("caption"),
            "dose_units": dose_units,
            "created": timezone.now().isoformat(),
            "last_updated": timezone.now().isoformat(),
            "rob_settings": AssessmentRiskOfBiasSerializer(self.assessment).data,
            "endpoints": [
                SerializerHelper.get_serialized(e, json=False) for e in self.get_endpoints(request)
            ],
            "studies": [
                SerializerHelper.get_serialized(s, json=False) for s in self.get_studies(request)
            ],
        }

    def get_rob_visual_type_display(self, value):
        rob_name = self.assessment.get_rob_name_display().lower()
        return value.replace("risk of bias", rob_name)

    def get_plotly_from_json(self) -> Figure:
        if self.visual_type != constants.VisualType.PLOTLY:
            raise ValueError("Incorrect visual type")
        try:
            return from_json(json.dumps(self.settings))
        except ValueError as err:
            raise ValueError(err) from err

    def _rob_data_qs(self, use_settings: bool = True) -> models.QuerySet:
        studies = self.get_studies()
        study_ids = (
            list(studies.values_list("id", flat=True))
            if isinstance(studies, models.QuerySet)
            else [study.id for study in studies]
        )
        settings = self.settings

        qs = RiskOfBiasScore.objects.filter(
            riskofbias__active=True,
            riskofbias__final=True,
            riskofbias__study__in=study_ids,
        )

        # use settings in read-only view; dont use when configuring plot
        if use_settings:
            qs = qs.filter(metric__in=settings["included_metrics"]).exclude(
                id__in=settings.get("excluded_score_ids", [])
            )
        return qs

    def data_df(self, use_settings: bool = True) -> pd.DataFrame:
        match self.visual_type:
            case constants.VisualType.ROB_BARCHART | constants.VisualType.ROB_HEATMAP:
                return self._rob_data_qs(use_settings=use_settings).df()
            case constants.VisualType.DATA_PIVOT_QUERY:
                return self._data_df_dpq().df
            case constants.VisualType.DATA_PIVOT_FILE:
                return self._data_df_dpf()
            case _:
                raise ValueError("Not supported for this visual type")

    def _data_df_dpq(self) -> FlatExport:
        def get_filterset(data, assessment, **kw):
            cls = prefilters.get_prefilter_cls(None, self.evidence_type, self.assessment)
            return cls(data=data, assessment=assessment, **kw)

        def get_queryset():
            prefilters = self.prefilters["prefilters"]
            preferred_units = self.prefilters["preferred_units"]
            fs = get_filterset(prefilters, self.assessment)
            form = fs.form
            fs.set_passthrough_options(form)
            qs = fs.qs
            if self.evidence_type == constants.StudyType.BIOASSAY and preferred_units:
                qs = qs.filter(animal_group__dosing_regime__doses__dose_units__in=preferred_units)
            return qs.order_by("id")

        def get_dataset_exporter(qs: models.QuerySet):
            export_style = self.prefilters["export_style"]
            preferred_units = self.prefilters["preferred_units"]
            if self.evidence_type == constants.StudyType.BIOASSAY:
                # select export class
                if export_style == constants.ExportStyle.EXPORT_GROUP:
                    ExportClass = EndpointGroupFlatDataPivot
                elif export_style == constants.ExportStyle.EXPORT_ENDPOINT:
                    ExportClass = EndpointFlatDataPivot
                else:
                    raise ValueError("Unknown export type")

                exporter = ExportClass(
                    qs,
                    assessment=self.assessment,
                    filename=f"{self.assessment}-animal-bioassay",
                    preferred_units=preferred_units,
                )

            elif self.evidence_type == constants.StudyType.EPI:
                if self.assessment.epi_version == EpiVersion.V1:
                    exporter = OutcomeDataPivot(
                        qs,
                        assessment=self.assessment,
                        filename=f"{self.assessment}-epi",
                    )
                else:
                    exporter = EpiFlatComplete(
                        qs,
                        assessment=self.assessment,
                        filename=f"{self.assessment}-epi",
                    )

            elif self.evidence_type == constants.StudyType.EPI_META:
                exporter = MetaResultFlatDataPivot(
                    qs,
                    assessment=self.assessment,
                    filename=f"{self.assessment}-epi",
                )

            elif self.evidence_type == constants.StudyType.IN_VITRO:
                # select export class
                if export_style == constants.ExportStyle.EXPORT_GROUP:
                    Exporter = ivexports.DataPivotEndpointGroup
                elif export_style == constants.ExportStyle.EXPORT_ENDPOINT:
                    Exporter = ivexports.DataPivotEndpoint
                else:
                    raise ValueError("Unknown export type")

                # generate export
                exporter = Exporter(
                    qs,
                    assessment=self.assessment,
                    filename=f"{self.assessment}-invitro",
                )

            elif self.evidence_type == constants.StudyType.ECO:
                exporter = EcoFlatComplete(
                    qs,
                    assessment=self.assessment,
                    filename=f"{self.assessment}-eco",
                )
            else:
                raise ValueError("Unknown export type")

            return exporter

        qs = get_queryset()
        exporter = get_dataset_exporter(qs)
        return exporter.build_export()

    def _data_df_dpf(self) -> pd.DataFrame:
        return self.dataset.get_latest_df()


reversion.register(SummaryTable)
reversion.register(Visual)
