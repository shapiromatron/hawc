from html.parser import HTMLParser
from typing import List

from docx import Document
from pydantic import BaseModel, conint, validator


class GenericTableCell(BaseModel):
    row: conint(ge=0)
    column: conint(ge=0)
    html_text: str

    def to_docx(self, par):
        parser = CellHTMLParser()
        return parser.feed(self.html_text, par)


class GenericTable(BaseModel):
    rows: conint(gt=0)
    columns: conint(gt=0)
    cells: List[GenericTableCell]

    @validator("cells")
    def validate_cells(cls, cells, values):
        for cell in cells:
            if cell.row >= values["rows"]:
                raise ValueError(
                    f"Cell row outside of row bounds: row index {cell.row}, number of rows {values['rows']}"
                )
            if cell.column >= values["columns"]:
                raise ValueError(
                    f"Cell column outside of column bounds: column index {cell.column}, number of columns {values['columns']}"
                )
        return cells

    def to_docx(self, docx: Document = None):
        return Document()
        if docx is None:
            docx = Document()
        table = docx.add_table(rows=self.rows, cols=self.columns)
        for cell in self.cells:
            table_cell = table.cell(cell.row, cell.column)
            cell_par = table_cell.add_paragraph()
            cell.to_docx(cell_par)
        return docx


class CellHTMLParser(HTMLParser):
    def feed(self, data, par):
        self.par = par
        super().feed(data)

    def handle_starttag(self, tag, attrs):
        pass

    def handle_data(self, data):
        self.par.add_run(data)

    def handle_endtag(self, tag):
        pass
