import docx
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches


class TableMaker:
    """
    Helper-object to build tables
    """

    def __init__(
        self,
        colWidths: list[float],
        styles: dict | None = None,
        numHeaders: int = 1,
        tblStyle: str | None = None,
        firstRowCaption: bool = True,
    ):
        """
        Args:
            colWidths (list[float]): Width of columns
            styles (dict, optional): Default styles for different fields.
                Keys include: ["title", "header", "body", "subheading"]. Defaults to None.
            numHeaders (int, optional): Number of headers. Defaults to 1.
            tblStyle (str, optional): Table style. Defaults to None.
            firstRowCaption (bool, optional): Whether first row is caption. Defaults to True.
        """
        self.colWidths = colWidths
        self.styles = styles or {}
        self.numHeaders = numHeaders
        self.tblStyle = tblStyle
        self.firstRowCaption = firstRowCaption
        self.cells = []
        self.rows = 0
        self.cols = len(self.colWidths)

    def render(self, doc):
        tbl = doc.add_table(rows=self.rows, cols=self.cols, style=self.tblStyle)

        # set column widths
        tbl.autofit = False
        for i, col in enumerate(tbl.columns):
            col.width = docx.shared.Inches(self.colWidths[i])

        # build cells
        for cell in self.cells:
            cell.render(tbl)

        # mark rows as headers to break on pages
        if self.numHeaders >= 1:
            for i in range(self.numHeaders):
                tblHeader = docx.oxml.parse_xml(
                    r"<w:tblHeader {} />".format(docx.oxml.ns.nsdecls("w"))
                )
                tbl.rows[i]._tr.get_or_add_trPr().append(tblHeader)

        # apply caption-style to the first cell in first-row
        if self.firstRowCaption:
            cell = tbl.cell(0, 0)

            cellPr = cell._tc.get_or_add_tcPr()
            cellPr.append(
                docx.oxml.parse_xml(
                    r'<w:tcBorders {} ><w:top w:val="nil"/><w:left w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'.format(
                        docx.oxml.ns.nsdecls("w")
                    )
                )
            )
            cellPr.append(
                docx.oxml.parse_xml(
                    r'<w:shd {} w:val="clear" w:color="auto" w:fill="auto"/>'.format(
                        docx.oxml.ns.nsdecls("w")
                    )
                )
            )

            # left-align text using justified (LEFT doesn't work)
            for p in cell.paragraphs:
                if p.style is None:
                    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY_LOW

        return tbl

    def get_style(self, style=None, default=None):
        # determine which style to use for a cell; otherwise return None
        if style:
            return style
        else:
            return self.styles.get(default, None)

    def _add_cell(self, cell):
        self.cells.append(cell)
        self.rows = max(self.rows, cell.row + 1)
        return cell

    def _get_width(self, col, colspan):
        if colspan:
            return sum(self.colWidths[col : col + colspan])
        else:
            return self.colWidths[col]

    def new_th(self, row, col, text, *arg, **kw):
        kw.update(
            row=row,
            col=col,
            width=self._get_width(col, kw.get("colspan")),
            style=self.get_style(style=kw.get("style")),
        )
        if kw["style"]:
            kw["text"] = text
        else:
            kw["runs"] = [TableMaker.new_run(text, b=True, newline=False)]
        return self._add_cell(CellMaker(**kw))

    def new_td_txt(self, row, col, text, *arg, **kw):
        kw.update(
            text=text,
            row=row,
            col=col,
            width=self._get_width(col, kw.get("colspan")),
            style=self.get_style(style=kw.get("style"), default="body"),
        )
        return self._add_cell(CellMaker(**kw))

    def new_td_run(self, row, col, runs, *arg, **kw):
        kw.update(
            runs=runs,
            row=row,
            col=col,
            width=self._get_width(col, kw.get("colspan")),
            style=self.get_style(style=kw.get("style"), default="body"),
        )
        return self._add_cell(CellMaker(**kw))

    @classmethod
    def new_run(cls, txt, b=False, i=False, newline=True, style=None):
        if newline:
            txt += "\n"
        return {"text": txt, "style": style, "bold": b, "italic": i}


class CellMaker:
    """
    Helper-object to build table-cells
    """

    def __init__(self, row: int, col: int, width: float, **kw):
        """
        Args:
            row (int)
            col (int)
            width (float)
            **kw: Optional keyword arguments, as follows:
                rowspan (int, optional)
                colspan (int, optional)
                style (str, optional)
                shade (str, optional)
                text (str, optional)
                runs (list[dict], optional)
                vertical (bool, optional)
                height (float, optional)
        """
        self.__dict__.update(row=row, col=col, width=width)
        self.__dict__.update(kw)

    def render(self, tbl):
        cellD = tbl.cell(self.row, self.col)
        p = cellD.paragraphs[0]

        # merge cells if needed
        rowspan = getattr(self, "rowspan", None)
        colspan = getattr(self, "colspan", None)
        if rowspan or colspan:
            rowIdx = self.row + (rowspan or 1) - 1
            colIdx = self.col + (colspan or 1) - 1
            cellD.merge(tbl.cell(rowIdx, colIdx))

        cellD.width = docx.shared.Inches(self.width)

        # add style
        style = getattr(self, "style", None)
        if style:
            p.style = style

        # add shading
        shade = getattr(self, "shade", None)
        if shade:
            tcPr = cellD._tc.get_or_add_tcPr()
            tcVAlign = OxmlElement("w:shd")
            tcVAlign.set(qn("w:fill"), shade)
            tcPr.append(tcVAlign)

        # change text rotation
        vertical = getattr(self, "vertical", False)
        if vertical:
            tcPr = cellD._tc.get_or_add_tcPr()
            tcVAlign = OxmlElement("w:textDirection")
            tcVAlign.set(qn("w:val"), "btLr")
            tcPr.append(tcVAlign)

        height = getattr(self, "height", False)
        if height:
            tcPr = cellD._tc.get_or_add_tcPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(Inches(float(height))))
            trHeight.set(qn("w:hRule"), "atLeast")
            tcPr.append(trHeight)

        # add content
        text = getattr(self, "text", None)
        if text:
            p.text = text

        runs = getattr(self, "runs", None)
        if runs:
            for runD in runs:
                run = p.add_run(runD["text"])
                run.bold = runD.get("bold", False)
                run.italic = runD.get("italic", False)
                run.style = runD.get("style", None)
