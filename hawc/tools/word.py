from docx import Document
from docx.enum import dml
from docx.opc import constants
from docx.oxml.shared import OxmlElement, qn


def write_setting_p(doc: Document, title: str, value: str):
    """Write a paragraph with a bolded title string followed by a value."""
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(value)


def add_url_hyperlink(paragraph, url: str, text: str):
    """Add URL hyperlink in docx report. Adapted from https://stackoverflow.com/a/47666747/906385

    Args:
        paragraph: a docx paragraph object
        url (str): The URL
        text (str): The text to display
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    run = paragraph.add_run()
    run._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    run.font.color.theme_color = dml.MSO_THEME_COLOR_INDEX.HYPERLINK
    run.font.underline = True
