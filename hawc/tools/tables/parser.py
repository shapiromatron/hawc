import re
import urllib.parse
from html.parser import HTMLParser

import docx


def tag_wrapper(text: str, tag: str, *args):
    if len(args) == 0:
        return f"<{tag}>{text}</{tag}>"
    arg_tag = args[-1]
    text = f"<{arg_tag}>{text}</{arg_tag}>"
    return tag_wrapper(text, tag, *args[:-1])


def strip_tags(text: str):
    return re.sub("<[^<]+?>", "", text)


def strip_enclosing_tag(text: str, tag: str):
    if text.startswith(f"<{tag}>") and text.endswith(f"</{tag}>"):
        return text[2 + len(tag) : -(3 + len(tag))]
    return text


def has_inner_text(text: str):
    # check for inner text by removing tags and whitespace
    return bool(strip_tags(text).strip())


def ul_wrapper(texts: list[str]):
    list_items = map(lambda text: tag_wrapper(text, "li"), texts)
    return f"<ul>{''.join(list_items)}</ul>"


def ol_wrapper(texts: list[str]):
    list_items = map(lambda text: tag_wrapper(text, "li"), texts)
    return f"<ol>{''.join(list_items)}</ol>"


def color_background(cell, color: str):
    """Add a cell background color to a table cell

    Args:
        cell: the cell to color
        color (str): A hex color string (eg., #4B9CD3)
    """
    cell_properties = cell._element.tcPr
    cell_shading = docx.oxml.shared.OxmlElement("w:shd")
    cell_shading.set(docx.oxml.shared.qn("w:fill"), color.lstrip("#"))
    cell_properties.append(cell_shading)


class QuillParser(HTMLParser):
    # Inline tags
    _inline_tags = ["strong", "em", "u", "a"]
    # Inline tags to their font attribute
    _tag_font_attr = {"strong": "bold", "em": "italic", "u": "underline"}
    # Cached inline tags
    _tags = set()

    # Paragraph tags
    _paragraph_tags = ["p", "h1", "h2", "li"]
    # Paragraph tags to their styles
    _tag_style = {
        "p": "Normal",
        "h1": "Heading 1",
        "h2": "Heading 2",
        "ol": "List Number",
        "ul": "List Bullet",
    }

    def __init__(self, *args, **kwargs):
        self.base_url = kwargs.pop("base_url", "")
        super().__init__(*args, **kwargs)

    def feed(self, data, block):
        # block is either a _Body or _Cell object
        self._block = block
        super().feed(data)

    def parse_url(self, url):
        # convert relative URL to absolute if a base URL is specified
        _url = urllib.parse.urlparse(url)
        if not (_url.scheme and _url.netloc) and self.base_url:
            url = urllib.parse.urljoin(self.base_url, url)
        return url

    def add_hyperlink(self):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = self._paragraph.part
        r_id = part.relate_to(
            self.parse_url(self._href),
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(
            docx.oxml.shared.qn("r:id"),
            r_id,
        )
        hyperlink.append(self._inline._element)
        self._paragraph._p.append(hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        self._inline.font.color.theme_color = docx.enum.dml.MSO_THEME_COLOR_INDEX.HYPERLINK
        self._inline.font.underline = True

    def handle_paragraph(self, tag):
        # We create a new paragraph and stylize it
        self._paragraph = self._block.add_paragraph()
        style = self._tag_style[self._list_type] if tag == "li" else self._tag_style[tag]
        self._paragraph.style = style
        # Then we create a new run for inline tags / data
        self._inline = self._paragraph.add_run()

    def handle_inline_start(self, tag, attrs):
        # A new inline tag means we have to create a new run
        self._inline = self._paragraph.add_run()
        if tag == "a":
            # We cache the href for hyperlink creation
            self._href = attrs[0][1]
        # We cache the tag so that we can later change the data font
        self._tags.add(tag)

    def handle_inline_end(self, tag):
        # A new run is created without the inline tag
        self._inline = self._paragraph.add_run()
        self._tags.remove(tag)

    def handle_starttag(self, tag, attrs):
        if tag == "ol" or tag == "ul":
            # We cache the list type to style list items
            self._list_type = tag
        elif tag in self._paragraph_tags:
            self.handle_paragraph(tag)
        elif tag in self._inline_tags:
            self.handle_inline_start(tag, attrs)

    def handle_data(self, data):
        self._inline.text += data
        # Apply the inline tag fonts/styles to the run with this data
        for tag in self._tags:
            if tag == "a":
                self.add_hyperlink()
            else:
                setattr(self._inline.font, self._tag_font_attr[tag], True)

    def handle_endtag(self, tag):
        if tag in self._inline_tags:
            self.handle_inline_end(tag)
