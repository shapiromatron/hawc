from enum import IntEnum
from typing import List

from docx import Document
from pydantic import BaseModel, conint, root_validator

from .parser import QuillParser


class CellEnum(IntEnum):
    CELL = 1
    EMPTY = 2
    SKIP = 3


class GenericTableCell(BaseModel):
    header: bool = False
    row: conint(ge=0)
    column: conint(ge=0)
    row_span: conint(ge=1) = 1
    col_span: conint(ge=1) = 1
    quill_text: str

    def __str__(self):
        return f"Cell (row={self.row}, column={self.column})"

    def row_order_index(self, columns):
        return self.row * columns + self.column

    def to_docx(self, block):
        parser = QuillParser()
        return parser.feed(self.quill_text, block)


class EvidenceIntegrationTable(BaseModel):
    rows: conint(gt=0)
    columns: conint(gt=0)
    cells: List[GenericTableCell]

    @root_validator(skip_on_failure=True)
    def validate_cells(cls, values):
        cells = values["cells"]
        table_cells = [[False] * values["columns"] for _ in range(values["rows"])]
        for cell in cells:
            try:
                for i in range(cell.row_span):
                    for j in range(cell.col_span):
                        if table_cells[cell.row + i][cell.column + j]:
                            raise ValueError(f"Cell overlap at {cell}")
                        table_cells[cell.row + i][cell.column + j] = True
            except IndexError:
                raise ValueError(f"{cell} outside of table bounds")
        cells.sort(key=lambda cell: cell.row_order_index(values["columns"]))
        return values

    def row_order_enums(self):
        # TODO move logic to JS for HTML table render(?)
        enums = [CellEnum.EMPTY] * (self.columns * self.rows)
        for cell in self.cells:
            cell_index = cell.row_order_index(self.rows)
            enums[cell_index] = CellEnum.CELL
            for i in range(1, cell.row_span):
                for j in range(1, cell.col_span):
                    index = cell_index + i + (j * self.rows)
                    enums[index] = CellEnum.SKIP

    def to_docx(self, docx: Document = None):
        if docx is None:
            docx = Document()
        table = docx.add_table(rows=self.rows, cols=self.columns)
        table.style = "Table Grid"
        for cell in self.cells:
            table_cell = table.cell(cell.row, cell.column)
            span_cell = table.cell(cell.row + cell.row_span - 1, cell.column + cell.col_span - 1)
            table_cell.merge(span_cell)
            # Remove default paragraph
            paragraph = table_cell.paragraphs[0]._element
            paragraph.getparent().remove(paragraph)
            cell.to_docx(table_cell)
        return docx

    @classmethod
    def build_default(cls):
        return cls.parse_raw(
            """
        {
            "rows": 2,
            "columns": 2,
            "cells": [
                {
                    "header": true,
                    "row": 0,
                    "column": 0,
                    "quill_text": "<p>A</p>"
                },
                {
                    "header": true,
                    "row": 0,
                    "column": 1,
                    "quill_text": "<p>B</p>"
                },
                {
                    "header": false,
                    "row": 1,
                    "column": 0,
                    "quill_text": "<p>C</p>"
                },
                {
                    "header": false,
                    "row": 1,
                    "column": 1,
                    "quill_text": "<p>D</p>"
                }
            ]
        }
        """
        )
