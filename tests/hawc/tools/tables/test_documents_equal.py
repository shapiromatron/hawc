from docx import Document

from . import documents_equal


def test_documents_equal_true():
    # documents with the same paragraphs are equal
    doc1 = Document()
    doc2 = Document()
    p1 = doc1.add_paragraph("paragraph")
    p2 = doc2.add_paragraph("paragraph")
    assert documents_equal(doc1, doc2)

    # documents with the same runs are equal
    r1 = p1.add_run("text")
    r2 = p2.add_run("text")
    assert documents_equal(doc1, doc2)

    # documents with the same styles are equal
    r1.bold = True
    r2.bold = True
    assert documents_equal(doc1, doc2)


def test_documents_equal_false():
    # documents with different paragraphs are unequal
    doc1 = Document()
    doc2 = Document()
    doc1.add_paragraph("paragraph")
    assert not documents_equal(doc1, doc2)

    # documents with different runs are unequal
    doc1 = Document()
    doc2 = Document()
    p1 = doc1.add_paragraph("paragraph")
    p2 = doc2.add_paragraph("paragraph")
    p1.add_run("text")
    assert not documents_equal(doc1, doc2)

    # documents with different styles are unequal
    doc1 = Document()
    doc2 = Document()
    p1 = doc1.add_paragraph("paragraph")
    p2 = doc2.add_paragraph("paragraph")
    r1 = p1.add_run("text")
    r1.bold = True
    p2.add_run("text")
    assert not documents_equal(doc1, doc2)
