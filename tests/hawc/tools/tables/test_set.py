from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from hawc.tools.tables.set.constants import AttributeChoices
from hawc.tools.tables.set.table import Column, StudyEvaluationTable

from . import documents_equal

FILE_PATH = Path(__file__).parent.absolute() / "data" / "set_report.docx"


class TestAttributeChoices:
    def test_free_html(self):
        attribute = AttributeChoices.FreeHtml
        col = Column.model_construct(attribute=attribute)

        # returns an empty cell; overridden by customizations
        series = pd.Series({"foo": "foo", "bar": "bar", "free_html": "free_html"})
        cell = col.get_cell(series)
        assert cell.quill_text == "<p></p>"
        series = pd.Series(dtype=float)
        cell = col.get_cell(series)
        assert cell.quill_text == "<p></p>"

    def test_rob(self):
        attribute = AttributeChoices.Rob
        col = Column.model_construct(attribute=attribute)

        # test with scores
        df = pd.DataFrame({"score_score": [1, 2]})
        cell = col.get_cell(df)
        assert cell.judgment == 1  # uses first value

        # test with no scores
        df = pd.DataFrame({"score_score": []})
        cell = col.get_cell(df)
        assert cell.judgment == -1

    def test_animal_group_doses(self):
        attribute = AttributeChoices.AnimalGroupDoses

        series = pd.Series(
            {"animal_group_doses": {"ppm": ["0, 1, 10", "1, 10, 100"], "mg/m3": ["1, 10, 100"]}}
        )

        # show all doses
        col = Column.model_construct(attribute=attribute, dose_unit="")
        cell = col.get_cell(series)
        assert cell.quill_text == "<p>0, 1, 10 ppm; 1, 10, 100 ppm; 1, 10, 100 mg/m3</p>"

        # show ppm doses
        col = Column.model_construct(attribute=attribute, dose_unit="ppm")
        cell = col.get_cell(series)
        assert cell.quill_text == "<p>0, 1, 10; 1, 10, 100</p>"

    def test_default(self):
        attribute = AttributeChoices.StudyShortCitation
        col = Column.model_construct(attribute=attribute)

        # valid
        series = pd.Series({"study_short_citation": "Foo et al.; Bar et al."})
        cell = col.get_cell(series)
        assert cell.quill_text == "<p>Foo et al.; Bar et al.</p>"


@pytest.mark.django_db
class TestStudyEvaluationTable:
    def test_docx(self, rewrite_data_files: bool):
        data = {
            "assessment_id": 1,
            "data_source": "ani",
            "published_only": False,
            "subheaders": [
                {"label": "Subheader above animal description", "start": 2, "length": 1}
            ],
            "columns": [
                {
                    "label": "Study citation",
                    "attribute": "study_short_citation",
                    "key": "1",
                    "width": 1,
                },
                {
                    "label": "Animal description",
                    "attribute": "animal_group_description",
                    "key": "2",
                    "width": 2,
                },
                {"label": "Rob score", "attribute": "rob", "metric_id": 1, "key": "3", "width": 3},
            ],
            "rows": [
                {"id": 1, "type": "study", "customized": []},
                {
                    "id": 1,
                    "type": "study",
                    "customized": [
                        {"key": "1", "html": "<p>Custom study name</p>"},
                        {"key": "2", "html": "<p>Custom animal description</p>"},
                        {"key": "3", "score_id": -1},
                    ],
                },
                {"id": -1, "type": "study", "customized": []},
            ],
        }
        table = StudyEvaluationTable.model_validate(data)
        document = table.to_docx()
        if rewrite_data_files:
            document.save(FILE_PATH)
        saved_document = Document(FILE_PATH)

        assert documents_equal(document, saved_document)
