from pathlib import Path

import pytest
from docx import Document

from hawc.tools.tables.generic import GenericTable

from . import documents_equal

FILE_PATH = Path(__file__).parent.absolute() / "data" / "generic_report.docx"


class TestGenericTable:
    def test_docx(self, rewrite_data_files: bool):
        document = Document()
        table = GenericTable.build_default()
        document = table.to_docx()
        if rewrite_data_files:
            document.save(FILE_PATH)
        saved_document = Document(FILE_PATH)

        assert documents_equal(document, saved_document)

    def test_good_parse(self):
        # all cell properties are present and valid
        obj = {
            "cells": [
                {
                    "row": 0,
                    "column": 0,
                    "col_span": 1,
                    "row_span": 1,
                    "quill_text": "<p>paragraph</p>",
                },
            ],
        }
        GenericTable.model_validate(obj)

    def test_bad_parse(self):
        # cells overlap
        obj = {
            "cells": [
                {
                    "row": 0,
                    "column": 0,
                    "col_span": 2,
                    "row_span": 1,
                    "quill_text": "<p>paragraph</p>",
                },
                {
                    "row": 0,
                    "column": 1,
                    "col_span": 1,
                    "row_span": 1,
                    "quill_text": "<p>paragraph</p>",
                },
            ],
        }
        with pytest.raises(ValueError, match="Cell overlap"):
            GenericTable.model_validate(obj)
