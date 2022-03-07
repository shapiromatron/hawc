from pathlib import Path

from docx import Document

from hawc.tools.tables.parser import QuillParser

from . import documents_equal

FILE_PATH = Path(__file__).parent.absolute() / "data" / "parsed_html.docx"


class TestQuillParser:
    def test_parser(self, rewrite_data_files: bool, quill_html):
        document = Document()
        parser = QuillParser()
        parser.feed(quill_html, document._body)
        if rewrite_data_files:
            document.save(FILE_PATH)
        saved_document = Document(FILE_PATH)

        assert documents_equal(document, saved_document)

    def test_parse_url(self):
        # no base url
        parser = QuillParser()
        assert parser.parse_url("/test/") == "/test/"
        assert parser.parse_url("https://def.com/test/") == "https://def.com/test/"

        # base url
        parser = QuillParser(base_url="https://abc.com")
        assert parser.parse_url("/test/") == "https://abc.com/test/"
        assert parser.parse_url("https://def.com/test/") == "https://def.com/test/"
