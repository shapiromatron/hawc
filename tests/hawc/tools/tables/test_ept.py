from pathlib import Path

from docx import Document

from hawc.tools.tables.ept import EvidenceProfileTable

from . import documents_equal

FILE_PATH = Path(__file__).parent.absolute() / "data" / "ept_report.docx"


class TestEvidenceProfileTable:
    def test_docx(self, rewrite_data_files: bool):
        document = Document()
        table = EvidenceProfileTable.build_default()
        document = table.to_docx()
        if rewrite_data_files:
            document.save(FILE_PATH)
        saved_document = Document(FILE_PATH)

        assert documents_equal(document, saved_document)
