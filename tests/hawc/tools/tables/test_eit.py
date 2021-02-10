from pathlib import Path

from docx import Document

from hawc.tools.tables.eit import EvidenceIntegrationTable

from . import documents_equal

DATA_PATH = Path(__file__).parent.absolute() / "data"


class TestEvidenceIntegrationTable:
    def test_docx(self):
        document = Document()
        table = EvidenceIntegrationTable.build_default()
        document = table.to_docx()
        saved_document = Document(DATA_PATH / "eit_report.docx")

        assert documents_equal(document, saved_document)
