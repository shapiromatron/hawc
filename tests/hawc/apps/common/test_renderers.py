import json
from io import BytesIO

import pandas as pd
import pytest
from django.urls import reverse
from docx import Document
from rest_framework.exceptions import MethodNotAllowed
from rest_framework.response import Response
from rest_framework.test import APIRequestFactory

from hawc.apps.common import renderers
from hawc.apps.common.helper import FlatExport, ReportExport
from hawc.tools.excel import get_writer, write_worksheet

from ..test_utils import get_client


@pytest.fixture
def basic_export():
    df = pd.DataFrame(data=[[1, 2], [3, 4]], columns=["a", "b"])
    return FlatExport(df=df, filename="fn")


class TestDocxRenderer:
    def test_success(self):
        docx = Document()
        docx.add_heading("Test123")
        export = ReportExport(docx=docx, filename="test")
        resp = Response()
        b = renderers.DocxRenderer().render(data=export, renderer_context={"response": resp})
        d2 = Document(BytesIO(b))
        assert resp["Content-Disposition"] == "attachment; filename=test.docx"
        assert d2.paragraphs[0].text == "Test123"

    def test_options(self):
        factory = APIRequestFactory()
        response = renderers.DocxRenderer().render(
            data={"dummy": "data"},
            renderer_context={"response": Response(), "request": factory.options("/")},
        )
        assert response == '{"dummy": "data"}'


class TestPandasBaseRenderer:
    def test_error(self):
        # test that our data-check passes; we use the Csv subclass to check
        with pytest.raises(ValueError):
            renderers.PandasCsvRenderer().render(
                data="not-a-dataframe", renderer_context={"response": Response()}
            )


class TestPandasCsvRenderer:
    def test_success(self, basic_export):
        response = renderers.PandasCsvRenderer().render(
            data=basic_export, renderer_context={"response": Response()}
        )
        assert response == "a,b\n1,2\n3,4\n"


class TestPandasTsvRenderer:
    def test_success(self, basic_export):
        response = renderers.PandasTsvRenderer().render(
            data=basic_export, renderer_context={"response": Response()}
        )
        assert response == "a\tb\n1\t2\n3\t4\n"


class TestPandasHtmlRenderer:
    def test_success(self, basic_export):
        response = renderers.PandasHtmlRenderer().render(
            data=basic_export, renderer_context={"response": Response()}
        )
        assert (
            response
            == '<table border="1" class="dataframe">\n  <thead>\n    <tr style="text-align: right;">\n      <th>a</th>\n      <th>b</th>\n    </tr>\n  </thead>\n  <tbody>\n    <tr>\n      <td>1</td>\n      <td>2</td>\n    </tr>\n    <tr>\n      <td>3</td>\n      <td>4</td>\n    </tr>\n  </tbody>\n</table>'
        )


class TestPandasJsonRenderer:
    def test_success(self, basic_export):
        response = renderers.PandasJsonRenderer().render(
            data=basic_export, renderer_context={"response": Response()}
        )
        assert json.loads(response) == [{"a": 1, "b": 2}, {"a": 3, "b": 4}]

        # test duplicate columns
        df = pd.DataFrame(data=[[1, 2], [3, 4]], columns=["a", "a"])
        export = FlatExport(df=df, filename="fn")
        response = renderers.PandasJsonRenderer().render(
            data=export, renderer_context={"response": Response()}
        )
        assert json.loads(response) == [{"a.1": 1, "a.2": 2}, {"a.1": 3, "a.2": 4}]


@pytest.mark.django_db
class TestBrowsableAPIRenderer:
    def test_success(self, db_keys):
        client = get_client("team", api=True)
        assert client.login(username="team@hawcproject.org", password="pw") is True
        url = reverse("animal:api:assessment-full-export", args=(db_keys.assessment_working,))
        resp = client.get(url + "?format=api")
        assert resp.status_code == 200


class TestXlsxRenderer:
    def test_success(self, basic_export):
        resp_obj = Response()
        assert "Content-Disposition" not in resp_obj

        response = renderers.PandasXlsxRenderer().render(
            data=basic_export, renderer_context={"response": resp_obj}
        )
        df2 = pd.read_excel(BytesIO(response))
        assert df2.to_dict(orient="records") == [{"a": 1, "b": 2}, {"a": 3, "b": 4}]
        assert resp_obj["Content-Disposition"] == "attachment; filename=fn.xlsx"

    def test_datetimes(self, basic_export):
        resp_obj = Response()
        assert "Content-Disposition" not in resp_obj

        # Datetimes with timezones are incompatible with Excel. Make sure that the renderer can handle DataFrames with timezoned datetimes.
        df = pd.DataFrame(
            data=[[1, "2000-01-01 01:00:00"], [2, "2005-12-31 02:10:00"]],
            columns=["count", "when"],
        )
        df = df.assign(when=df.when.astype("datetime64[ns]"))

        def check_is_close(df1: pd.DataFrame, df2: pd.DataFrame) -> bool:
            return all(
                [
                    df1["count"].equals(df2["count"]),
                    bool(((df2["when"] - df1["when"]).dt.total_seconds().abs() < 0.1).all()),
                ]
            )

        # naive datetime
        response = renderers.PandasXlsxRenderer().render(
            data=FlatExport(df=df, filename="fn"), renderer_context={"response": Response()}
        )
        df2 = pd.read_excel(BytesIO(response))
        assert check_is_close(df, df2) is True

        # with timezone
        df.loc[:, "when"] = df.when.dt.tz_localize(tz="US/Eastern")
        response = renderers.PandasXlsxRenderer().render(
            data=FlatExport(df=df, filename="fn"), renderer_context={"response": Response()}
        )
        df2 = pd.read_excel(BytesIO(response))

        # this should be incorrect initially without a timezone
        with pytest.raises(TypeError):
            check_is_close(df, df2)

        # with appropriate cast, success!
        df2.loc[:, "when"] = df2.when.astype("datetime64[ns]").dt.tz_localize(tz="US/Eastern")
        assert check_is_close(df, df2) is True

    def test_invalid_chars(self):
        resp_obj = Response()
        assert "Content-Disposition" not in resp_obj

        # ensure IllegalCharacterError are properly caught
        df3 = pd.DataFrame(data=[["test-\x02-test"]], columns=["test"])
        response = renderers.PandasXlsxRenderer().render(
            data=FlatExport(df3, "name"), renderer_context={"response": resp_obj}
        )
        df2 = pd.read_excel(BytesIO(response))
        assert df2.to_dict(orient="records") == [{"test": "test--test"}]
        assert resp_obj["Content-Disposition"] == "attachment; filename=name.xlsx"

    def test_response_error(self):
        request_exception = MethodNotAllowed(method="POST")
        response = Response(
            data={"detail": str(request_exception)}, status=request_exception.status_code
        )
        # these properties are usually set by the view;
        # we will set them manually
        response.accepted_renderer = renderers.PandasXlsxRenderer()
        response.accepted_media_type = response.accepted_renderer.media_type
        response.renderer_context = {"response": response}
        response.render()
        # the rendered content should be a binary JSON with the exception details
        assert response.rendered_content == b'{"detail": "Method \\"POST\\" not allowed."}'

    def test_options_request(self):
        """
        Make sure that sending an OPTIONS to an xlsx-style export doesn't result in server error.

        This test was added based on security scan; please don't remove.
        """
        # We will pass in an OPTIONS request to the renderer context
        factory = APIRequestFactory()
        request = factory.options("/path")
        # The response data from an OPTIONS request is type dict
        data = {"dummy": "data"}
        response = renderers.PandasXlsxRenderer().render(
            data=data, renderer_context={"response": Response(), "request": request}
        )
        # The renderered response should be a JSON string of the passed in data
        assert response == '{"dummy": "data"}'


class TestXlsxBinaryRenderer:
    def test_success(self):
        resp_obj = Response()
        assert "Content-Disposition" not in resp_obj

        f, writer = get_writer()
        with writer:
            write_worksheet(writer, "foo", pd.DataFrame(data=[[1, 2], [3, 4]], columns=["a", "b"]))

        response = renderers.XlsxBinaryRenderer().render(
            data=renderers.BinaryXlsxDataFormat(bytes=f, filename="fn"),
            renderer_context={"response": resp_obj},
        )
        df2 = pd.read_excel(BytesIO(response))
        assert df2.to_dict(orient="records") == [{"a": 1, "b": 2}, {"a": 3, "b": 4}]
        assert resp_obj["Content-Disposition"] == "attachment; filename=fn.xlsx"

    def test_options(self):
        factory = APIRequestFactory()
        response = renderers.XlsxBinaryRenderer().render(
            data={"dummy": "data"},
            renderer_context={"response": Response(), "request": factory.options("/")},
        )
        assert response == '{"dummy": "data"}'

    def test_bad_request(self):
        factory = APIRequestFactory()
        with pytest.raises(ValueError):
            renderers.XlsxBinaryRenderer().render(
                data=123, renderer_context={"response": Response(), "request": factory.options("/")}
            )
