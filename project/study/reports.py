from django.utils.html import strip_tags

from docx.shared import Inches

from utils.helper import DOCXReport

from . import models


def build_header_cell(row, col, width, text, colspan=1):
    return {
        "row": row,
        "col": col,
        "width": width,
        "colspan": colspan,
        "runs": [{ "text": text, "bold": True, "italic": False }]}

def build_caption(study):
    cell = build_header_cell(0, 0, 7.5, "", colspan=3)
    cell["runs"] = [
        { "text": u'{0}: '.format(study["short_citation"]), "bold": True, "italic": False },
        { "text": study["full_citation"], "bold": False, "italic": False }
    ]
    return cell


class SQDOCXReport(DOCXReport):

    def build_ROB_table(self, study):

        # build header rows
        cells = [
            build_caption(study),
            build_header_cell(1, 0, 2.8, r'Risk of Bias'),
            build_header_cell(1, 1, 0.7, r'Rating'),
            build_header_cell(1, 2, 4.0, r'Rating Rationale'),
        ]

        # build RoB rows
        rows = 2
        for d in study["qualities"]:
            cells.append({"row": rows, "col": 0, "text": d["metric"]["metric"]})
            cells.append({"row": rows, "col": 1, "text": d["score_symbol"], "shade": d["score_shade"]})
            cells.append({"row": rows, "col": 2, "text": strip_tags(d["notes"])})
            rows += 1

        self.build_table(rows, 3, cells)

    def build_ROB_legend(self):
        cells = [
           {"row": 0, "col": 0, "width": 1.5, "runs": [
            {"text": "Risk of Bias Key:", "bold": True, "italic": False}
           ]},
           {"row": 0, "col": 1, "width": 1.0, "text": "Definitely Low"},
           {"row": 0, "col": 2, "width": 0.4, "text": "++",
            "shade": models.StudyQuality.SCORE_SHADES[4]},
           {"row": 0, "col": 3, "width": 1.0, "text": "Probably Low"},
           {"row": 0, "col": 4, "width": 0.4, "text": "+",
            "shade": models.StudyQuality.SCORE_SHADES[3]},
           {"row": 0, "col": 5, "width": 1.0, "text": "Probably High"},
           {"row": 0, "col": 6, "width": 0.4, "text": "-",
            "shade": models.StudyQuality.SCORE_SHADES[2]},
           {"row": 0, "col": 7, "width": 0.4, "text": "NR",
            "shade": models.StudyQuality.SCORE_SHADES[0]},
           {"row": 0, "col": 8, "width": 1.0, "text": "Definitely High"},
           {"row": 0, "col": 9, "width": 0.4, "text": "--",
            "shade": models.StudyQuality.SCORE_SHADES[1]},
        ]
        self.build_table(1, len(cells), cells)

    def create_context(self):
        doc = self.doc
        d = self.context

        # make landscape
        section = doc.sections[-1]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        # title
        txt = "Risk of bias report: {0}".format(d['assessment']['name'])
        doc.add_heading(txt, 0)

        # loop through each study
        for study in d['studies']:
            doc.add_heading(study['short_citation'], 1)
            if len(study["qualities"])>0:
                self.build_ROB_table(study)
                doc.add_paragraph()
                self.build_ROB_legend()
                doc.add_page_break()
            else:
                doc.add_paragraph("No risk-of-bias details available for this study.")
