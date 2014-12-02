from json import dumps, loads
import logging
import os
import random
import string

from django.db import models
from django.conf import settings
from django.core.files import File
from django.core.serializers.json import DjangoJSONEncoder
from django.core.urlresolvers import reverse

from docx.shared import Inches

from bmds.bmds import BMDS
from bmds.output_parser import BMD_output_parser
from utils.executable_sheller import run_process
from utils.helper import HAWCdocx


BMDS_CHOICES = (
    ('2.30', 'Version 2.30'),
    ('2.31', 'Version 2.31'),
    ('2.40', 'Version 2.40'),
)

LOGIC_BIN_CHOICES = (
    (0, 'Warning (no change)'),
    (1, 'Questionable'),
    (2, 'Not Viable'),
)


class BMD_session(models.Model):
    endpoint = models.ForeignKey('animal.Endpoint',
                                 related_name='BMD_session', null=True)
    dose_units = models.ForeignKey('animal.DoseUnits')
    BMDS_version = models.CharField(max_length=10, choices=BMDS_CHOICES)
    created = models.DateTimeField(auto_now_add=True)
    last_updated = models.DateTimeField(auto_now=True)
    bmrs = models.TextField(blank=True)  # list of bmrs, json encoded
    selected_model = models.OneToOneField('BMD_model_run', blank=True, null=True, related_name='selected')
    notes = models.TextField()

    class Meta:
        get_latest_by = "last_updated"

    @classmethod
    def get_template(cls, assessment, data_type, json=False):
        """
        Return a template that can be used for a BMD modeling session.
        Requires an assessment to determine the correct BMD modeling version
        and an endpoint.data_type to return the proper collection of models to
        be used.
        """
        bmds_version = assessment.BMD_Settings.BMDS_version
        bmds = BMDS.versions[bmds_version]
        data = {'session': {}, 'models': []}
        data['session']['bmrs'] = bmds.template_bmrs[data_type]
        for m in bmds.template_models[data_type]:
            run_instance = bmds.models[data_type][m]()
            model_data = BMD_model_run.get_model_template(m, run_instance)
            data['models'].append(model_data)
        return dumps(data, cls=DjangoJSONEncoder) if json else data

    def save(self, *args, **kwargs):
        super(BMD_session, self).save(*args, **kwargs)
        # circular module import errors force us to call using this method style
        self.endpoint.__class__.d_response_delete_cache([self.endpoint.pk])

    def get_selected_model(self, json_encode=True):
        if self.selected_model:
            return self.selected_model.webpage_return(json=json_encode)
        else:
            return dumps(None)

    def docx_print(self, report, heading_level):
        """ Word report format for bmd printing """

        # define content
        title = 'Benchmark Dose Modeling'
        paras = (
            'BMDS software version: {0} (HAWC build)'.format(self.BMDS_version),
            'Created: {0}'.format(HAWCdocx.to_date_string(self.created)),
            'Last Updated: {0}'.format(HAWCdocx.to_date_string(self.last_updated)),
            )

        # save BMD summary table
        models = list(BMD_model_run.objects.filter(BMD_session=self.pk))
        rows = [['Model name', 'Global p-value', 'AIC', 'BMD', 'BMDL', 'Residual of Interest']]
        for model in models:
            rows.append(model.docx_print_table_row())

        if self.selected_model:
            models.pop(models.index(self.selected_model))

        # print document
        report.doc.add_heading(title, level=heading_level)
        for para in paras:
            report.doc.add_paragraph(para)

        tbl = report.doc.add_table(len(rows), len(rows[0]))
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                tbl.cell(i, j).text = val

        report.doc.add_page_break()

        # print selected model
        if self.selected_model:
            self.selected_model.docx_print(report, heading_level=heading_level+1, selected=True)
        else:
            p = report.doc.add_paragraph('No model was selected as a best-fitting model.')
            p.italic = True

        # print remaining models
        for model in models:
            model.docx_print(report, heading_level=heading_level+1, selected=False)

    def webpage_return(self, json=False):
        """
        Return an object specifying all the settings, configurations, and
        outputs from a previously run BMDS session.
        """
        try:
            session = self.selected_model.pk
        except:
            session = -1
        data = {
            'session': {'BMDS_version': self.BMDS_version,
                        'pk': self.pk,
                        'selected_model': session,
                        'notes': self.notes,
                        'created': self.created,
                        'last_updated': self.last_updated,
                        'bmrs': loads(self.bmrs)},
            'models': [],
        }
        runs = BMD_model_run.objects.filter(BMD_session=self.pk).order_by('option_id', 'bmr_id')
        for run in runs:
            data['models'].append(run.webpage_return())
        return dumps(data, cls=DjangoJSONEncoder) if json else data

    def run_session(self, endpoint, model_settings):
        """
        Run a BMD session, given an endpoint ID and a session. Returns
        a the results from a bmd session if successful.
        """

        try:
            dataset = endpoint.d_response(json_encode=False, dose_pk=model_settings['dose_units_id'])
            if 'warnings' in dataset:
                raise Exception('Dose units could not be resolved')
        except Exception:
            return False

        BMDS_version = endpoint.get_assessment().BMD_Settings.BMDS_version
        bmds = BMDS.versions[BMDS_version]()
        for option_id, option in enumerate(model_settings['options']):
            logging.debug('Creating ' + option['model_name'])
            #cycle through BMR to check if BMRs are valid
            for bmr_id, bmr in enumerate(model_settings['bmrs']):
                logging.debug('Testing ' + endpoint.data_type + ' ' + option['model_name'])
                run_instance = bmds.models[endpoint.data_type][option['model_name']]()
                # run instance in model/BMR type relationship is valid
                if run_instance.valid_bmr(bmr):
                    #create default model and override with user-settings
                    logging.debug('Updating options')
                    run_instance.update_model(option['override'],
                                              option['override_text'], bmr)

                    logging.debug('Creating new model instance')
                    run = BMD_model_run(model_name=option['model_name'],
                                        BMD_session=self,
                                        option_defaults=run_instance.build_defaults(json=True),
                                        option_override=dumps(option['override']),
                                        option_override_text=dumps(option['override_text']),
                                        option_id=option_id,
                                        bmr_id=bmr_id)
                    try:
                        logging.debug('Running ' + option['model_name'])
                        run.run_model(bmds, run_instance, dataset)
                    except:
                        logging.debug('Runtime error occurred for ' + option['model_name'])
                        run.runtime_error = True
                        run.output_text = 'An error occurred when running this model.'

                    logging.debug('Adding outputs ' + option['model_name'])
                    run.save()
        return True

    def get_assessment(self):
        return self.endpoint.get_assessment()


class BMD_model_run(models.Model):  # todo: get rid of these blank=True fields?
    model_name = models.CharField(max_length=25, blank=False)
    BMD_session = models.ForeignKey('BMD_session')
    option_defaults = models.TextField()  # json text
    option_override = models.TextField(blank=True)  # json text
    option_override_text = models.TextField(blank=True)  # json text
    output_text = models.TextField(blank=True)  # bmds raw output file
    outputs = models.TextField(blank=True)  # json text
    d3_plotting = models.TextField(blank=True)
    runtime_error = models.BooleanField(default=False)
    plot = models.ImageField(upload_to='bmds_plot', blank=True, null=True)
    option_id = models.PositiveSmallIntegerField()
    bmr_id = models.PositiveSmallIntegerField()
    override = models.PositiveSmallIntegerField(default=99)
    override_text = models.TextField(default="")

    @classmethod
    def get_model_template(cls, model_name, run_instance):
        """
        Return a model template for the selected BMD model.
        """
        run = BMD_model_run(model_name=model_name,
                            option_defaults=run_instance.build_defaults(json=True),
                            option_override="{}",
                            option_override_text="[]",
                            runtime_error=True,
                            id='',
                            bmr_id='',
                            option_id='')
        return run.webpage_return()

    @staticmethod
    def delete_files(filelist):
        """attempt to kill each file in filelst"""
        for f in filelist:
            try:
                os.remove(f)
            except:
                pass

    def get_assessment(self):
        return self.BMD_session.endpoint.get_assessment()

    def __unicode__(self):
        return self.model_name

    def build_d_file(self, endpoint):
        """
        Reconstruct a d-file based on the current option settings for the
        endpoint of interest.
        """
        bmds = BMDS.versions[self.BMD_session.BMDS_version]()
        run_instance = bmds.models[endpoint.data_type][self.model_name]()
        run_instance.update_model(loads(self.option_override),
                                  loads(self.option_override_text),
                                  self.get_bmr_dict())
        return run_instance.dfile_print(endpoint.d_response(json_encode=False))

    def get_bmr_dict(self):
        """
        Return a BMR dictionary for this model's BMR.
        """
        bmrs = loads(self.BMD_session.bmrs)
        return bmrs[self.bmr_id]

    def get_bmr_text(self, short_text=False):
        """
        Print a descriptive name for the bmr text used. Can either return a
        short_texxt name (for tables), or a longer more descriptive name.
        """
        bmr = self.get_bmr_dict()
        table_type_dict = {
            'Extra': {'short': r'%', 'long': r'% Extra Risk'},
            'Added': {'short': r'% AR', 'long': r'% Added Risk'},
            'Abs. Dev.': {'short': r' AD', 'long': r' Absolute Deviation(s)'},
            'Std. Dev.': {'short': r' SD', 'long': r' Standard Deviation(s)'},
            'Rel. Dev.': {'short': r'% RD', 'long': r'% Relative Deviation(s)'},
            'Point': {'short': r' Pt', 'long': r' Point'},
        }
        str_type = 'short' if short_text else 'long'
        if bmr['type'] in ['Extra', 'Added', 'Rel. Dev.']:
            return str(bmr['value'] * 100.) + table_type_dict[bmr['type']][str_type]
        elif bmr['type'] in ['Abs. Dev.', 'Std. Dev.', 'Point']:
            return str(bmr['value']) + table_type_dict[bmr['type']][str_type]
        else:
            return str(bmr['value']) + ' ' + bmr['type']

    def webpage_return(self, json=False):
        try:
            bmds_plot_url = self.plot.url
        except:
            bmds_plot_url = ''

        outputs = {'model_name': self.model_name,
                   'output_text': self.output_text,
                   'option_defaults': loads(self.option_defaults),
                   'option_override': loads(self.option_override),
                   'option_override_text': loads(self.option_override_text),
                   'id': self.id,
                   'bmds_plot_url': bmds_plot_url,
                   'option_id': self.option_id,
                   'bmr_id': self.bmr_id,
                   'override': self.override,
                   'override_text': self.override_text}
        if hasattr(self, 'BMD_session'):
            outputs['dose_units_id'] = self.BMD_session.dose_units.pk

        if self.runtime_error:
            outputs['plotting'] = 'error'
            outputs['outputs'] = 'error'
        else:
            outputs['plotting'] = loads(self.d3_plotting)
            outputs['outputs'] = loads(self.outputs)
        return dumps(outputs, cls=DjangoJSONEncoder) if json else outputs

    def d3_plot_info(self, model_class, dataset):
        p = {}
        outs = loads(self.outputs)
        for parameter in model_class.js_parameters:
            try:
                p[parameter] = outs['parameters'][parameter]['estimate']
            except:
                p[parameter] = 0.
        if 'dataset_increasing' in dataset:  # this is required for exponential M2/M3
            if dataset['dataset_increasing']:
                p['sign'] = 1.
            else:
                p['sign'] = -1.
        return {'formula': model_class.js_formula, 'parameters': p}

    def run_model(self, BMDS, model_class, d_response):
        """
        Given
            BMDS - BMDS version object
            model_class - BMDS model object
            d_response - endpoint d_response object

        Run the model and import results.
        """

        #Exit-early if number of dose-groups less than required minimum
        d_response['numDG'] = len(d_response['dr'])
        if d_response['numDG'] < model_class.minimum_DG:
            self.plot = None
            self.output_text = 'Dose groups less than minimum required to run model.'
            self.runtime_error = True
            return False

        # build d_file, save d-file options to object
        ds = model_class.dfile_print(d_response)

        # run BMDS
        outputs = self.execute_bmds(BMDS, model_class, ds)

        if outputs.get('image'):
            self.plot.save(outputs['image_fn'], outputs['image'])

        self.output_text = outputs.get('output_text')
        self.outputs = dumps(outputs.get('outputs'))
        self.d3_plotting = dumps(self.d3_plot_info(model_class, d_response))

    def execute_bmds(self, BMDS, model_class, d_file, create_image=True):
        """
        Run series of executables and return a dictionary of results
        """
        outputs = {}

        # get random filename
        temp_fn = ''.join(random.choice(string.ascii_lowercase) for x in range(16))
        f_in = os.path.join(BMDS.temp_path, temp_fn) + '.(d)'
        #exponential models
        prefix = getattr(model_class, 'output_prefix', "")
        temp_fn = prefix + temp_fn
        f_out = os.path.join(BMDS.temp_path, temp_fn) + '.out'

        with open(f_in, 'w') as f:
            f.write(d_file)

        # run BMDS
        logging.debug('running BMDS model...')
        exe = os.path.join(BMDS.model_path, model_class.exe + settings.BMD_EXTENSION)
        run_process([exe, f_in], 10).Run()

        if create_image:
            outputs['image'] = self.execute_bmds_figure(BMDS, model_class, temp_fn)

        # import outputs
        logging.debug('importing BMDS results...')
        output_text=""
        if os.path.exists(f_out):
            with open(f_out, 'r') as f:
                output_text = f.read()

        if model_class.model_name == 'Exponential':
            dtype = 'E'
        else:
            dtype = model_class.dtype
        o = BMD_output_parser(output_text, dtype, self.model_name)

        outputs['image_fn'] = temp_fn + '.emf'
        outputs['output_text'] = output_text
        outputs['outputs'] = o.output

        BMD_model_run.delete_files(([f_in, f_out]))
        return outputs

    def execute_bmds_figure(self, BMDS, model_class, temp_fn):
        """
        Execute BMDS executable to generate BMD figure.
        """
        image = None
        f_002 = os.path.join(BMDS.temp_path, temp_fn) + '.002'
        f_plt = os.path.join(BMDS.temp_path, temp_fn) + '.plt'
        f_emf = os.path.join(BMDS.temp_path, temp_fn) + '.emf'

        # Attempt to generate wgnuplot output images- an image processing
        # failure will not result in an overall error in the BMD modeling
        try:
            # Create a .plt file from a .002 file
            exe_plt = os.path.join(BMDS.model_path,
                                   model_class.exe_plot + settings.BMD_EXTENSION)
            run_process([exe_plt, f_002], 10).Run()

            # Revise settings of .plt file to export as emf image instead
            #  of appear in the GUI
            # todo: recompile BMDS source to do this automatically
            plotfile = open(f_plt, 'r')
            plot_text = plotfile.read()
            plotfile.close()
            plot_text = plot_text.replace("set terminal %s \n" % (settings.BMD_SHELL),
                                          "set terminal emf \nset output '%s'\n" % (f_emf))
            plot_text = plot_text.replace("set terminal %s\n" % (settings.BMD_SHELL),
                                          "set terminal emf \nset output '%s'\n" % (f_emf))
            plotfile = open(f_plt, 'w')
            plotfile.write(plot_text)
            plotfile.close()

            # Call gnuplot to generate image
            gnu_exe = os.path.join(BMDS.model_path,
                                   settings.BMD_PLOT + settings.BMD_EXTENSION)
            run_process([gnu_exe, f_plt], 10).Run()

            # Load plot into memory
            if os.path.exists(f_emf) and (os.path.getsize > 0):
                image = File(open(f_emf, 'rb'))
                logging.debug("gnuplot plot generation successful")
            else:
                logging.debug("gnuplot plot generation not successful")

        except:
            logging.debug("gnuplot plot generation not successful")

        BMD_model_run.delete_files(([f_002, f_plt, f_emf]))
        return image

    def docx_print(self, report, heading_level, selected):
        """
        Print outputs of self to Word document
        """
        title = '{0} Output Text'.format(self.model_name)
        if selected:
            title = 'Selected Best-Fitting Model: ' + title

        # image-file
        try:
            fn = self.plot.path
            # ajs todo: implement EMF in python-docx
            report.doc.add_picture(fn, width=Inches(4))
        except:
            report.doc.add_paragraph("No BMDS output image available.")

        # text
        report.doc.add_heading(title, level=heading_level)
        for line in self.output_text.split('\n'):
            report.doc.add_paragraph(text=line, style='RawOutput')
        report.doc.add_page_break()

    def docx_print_table_row(self):
        """
        Return list of string values to insert into summary BMD table.
        """
        output = {}
        try:
            output = loads(self.outputs)
        except ValueError:
            pass

        row = [self.model_name]
        for col in ('p_value4', 'AIC', 'BMD', 'BMDL', 'residual_of_interest'):
            row.append(str(output.get(col, '')))
        return row

    def summary_data(self):
        """
        Return a dictionary of key summary data useful for displaying in templates.
        Typically used for selected model in a BMDS session.
        """
        outputs = loads(self.outputs)
        summary_data = {'model_name': self.model_name}
        summary_data.update(outputs)
        return summary_data


class BMD_Assessment_Settings(models.Model):
    assessment = models.OneToOneField('assessment.Assessment',
                                      related_name='BMD_Settings')
    BMDS_version = models.CharField(max_length=10, choices=BMDS_CHOICES,
                                    default=max(BMDS_CHOICES)[0])
    created = models.DateTimeField(auto_now_add=True)
    last_updated = models.DateTimeField(auto_now=True)

    class Meta:
        verbose_name_plural = "BMD assessment settings"

    def __unicode__(self):
        return self.assessment.__unicode__() + ' BMD settings'

    def get_absolute_url(self):
        return reverse('bmd:assess_settings_detail', args=[self.assessment.pk])

    def get_assessment(self):
        return self.assessment


class LogicField(models.Model):
    """
    BMD decision-logic selections. Specific to an entire BMD assessment
    """
    assessment = models.ForeignKey('assessment.Assessment', related_name='BMD_Logic_Fields', editable=False)
    created = models.DateTimeField(auto_now_add=True)
    last_updated = models.DateTimeField(auto_now=True)
    logic_id = models.PositiveSmallIntegerField(editable=False)
    name = models.CharField(max_length=30, editable=False)
    function_name = models.CharField(max_length=25, editable=False)
    description = models.TextField(editable=False)
    failure_bin = models.PositiveSmallIntegerField(
        choices=LOGIC_BIN_CHOICES, blank=False,
        help_text="If the test fails, select the model-bin should the model be placed into.")
    threshold = models.DecimalField(
        max_digits=5, decimal_places=2, null=True, blank=True,
        help_text="If a threshold is required for the test, threshold can be specified to non-default.")
    continuous_on = models.BooleanField(default=True, verbose_name="Continuous Datasets")
    dichotomous_on = models.BooleanField(default=True, verbose_name="Dichotomous Datasets")
    cancer_dichotomous_on = models.BooleanField(default=True, verbose_name="Cancer Dichotomous Datasets")

    class Meta:
        ordering = ['logic_id']

    def __unicode__(self):
        return self.description

    def webpage_return(self, endpoint_data_type, json=False):
        outputs = {'last_updated': self.last_updated,
                   'logic_id': self.logic_id,
                   'name': self.name,
                   'function_name': self.function_name,
                   'description': self.description,
                   'failure_bin': self.failure_bin,
                   'threshold': self.threshold,
                   'test_on': self.datatype_inclusion(endpoint_data_type)}
        return dumps(outputs, cls=DjangoJSONEncoder) if json else outputs

    def datatype_inclusion(self, endpoint_data_type):
        crosswalk = {'C': self.continuous_on,
                     'D': self.dichotomous_on,
                     'DC': self.cancer_dichotomous_on}
        try:
            return crosswalk[endpoint_data_type]
        except:
            return False

    def get_assessment(self):
        return self.assessment

    @classmethod
    def website_list_return(cls, objs, endpoint_data_type, json=False):
        logics = []
        for obj in objs:
            logics.append(obj.webpage_return(endpoint_data_type, False))
        return dumps(logics, cls=DjangoJSONEncoder) if json else logics

    @classmethod
    def build_default(cls, id):
        """
        Constructor to build a single BMD decision logic setting.
        """
        if (int(id) < len(LogicField.default)):
            return LogicField(**LogicField.default[id])
        else:
            return None

    @classmethod
    def build_defaults(cls, assessment):
        """
        Constructor to build all default BMD decision logic settings.
        """
        fields = []
        for i in xrange(len(LogicField.default)):
            f = LogicField.build_default(i)
            f.assessment = assessment
            f.logic_id = i
            fields.append(f)
        LogicField.objects.bulk_create(fields)


LogicField.default = [
    {'failure_bin': 2,
     'description': 'A BMD value was successfully calculated.',
     'function_name': 'exists_bmd',
     'name': 'BMD',
     'threshold': None,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 2,
     'description': 'A BMDL value was successfully calculated.',
     'function_name': 'exists_bmdl',
     'name': 'BMDL',
     'threshold': None,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 0,
     'description': 'A BMDU value was successfully calculated.',
     'function_name': 'exists_bmdu',
     'name': 'BMDU',
     'threshold': None,
     'continuous_on': False,
     'dichotomous_on': False,
     'cancer_dichotomous_on': True},
    {'failure_bin': 2,
     'description': 'An AIC value was successfully calculated.',
     'function_name': 'exists_aic',
     'name': 'AIC',
     'threshold': None,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 1,
     'description': 'The correct variance model was used (based on p-test #2).',
     'function_name': 'variance_type',
     'name': 'Variance Type',
     'threshold': 0.1,
     'continuous_on': True,
     'dichotomous_on': False,
     'cancer_dichotomous_on': False},
    {'failure_bin': 1,
     'description': 'The variance model appropriately captures the variance of '
                    'the dataset being modeled (based on p-test #3).',
     'function_name': 'variance_fit',
     'name': 'Variance Fit',
     'threshold': 0.1,
     'continuous_on': True,
     'dichotomous_on': False,
     'cancer_dichotomous_on': False},
    {'failure_bin': 1,
     'description': 'The mean model is appropriately modeling the dataset, '
                    'based on the global-goodness-of-fit (p-test #4).',
     'function_name': 'ggof',
     'name': 'GGOF',
     'threshold': 0.1,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': False},
    {'failure_bin': 1,
     'description': 'The mean model is appropriately modeling the dataset, '
                    'based on the global-goodness-of-fit (p-test #4).',
     'function_name': 'ggof',
     'name': 'GGOF (Cancer)',
     'threshold': 0.05,
     'continuous_on': False,
     'dichotomous_on': False,
     'cancer_dichotomous_on': True},
    {'failure_bin': 1,
     'description': 'The ratio between the BMD and BMDL values is large, '
                    'indicating a large spread of estimates in the range of '
                    'interest (warning).',
     'function_name': 'bmd_ratio',
     'name': 'BMD/BMDL (serious)',
     'threshold': 20,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 0,
     'description': 'The ratio between the BMD and BMDL values is large, '
                    'indicating a large spread of estimates in the range of '
                    'interest (caution).',
     'function_name': 'bmd_ratio',
     'name': 'BMD/BMDL (warning)',
     'threshold': 5,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 1,
     'description': 'The absolute value of the residual at the dose closest to '
                    'the BMD is large, which means the model may not be '
                    'estimating well in this range.',
     'function_name': 'residual',
     'name': 'Residual of Interest',
     'threshold': 2,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 0,
     'description': 'The model output file included additional warnings.',
     'function_name': 'warnings',
     'name': 'Warnings',
     'threshold': None,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 0,
     'description': 'The BMD estimate is higher than the maximum dose modeled, '
                    'thus the model is extrapolating a result.',
     'function_name': 'high_bmd',
     'name': 'BMD higher',
     'threshold': 1,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 1,
     'description': 'The BMDL estimate is higher than the maximum dose modeled, '
                    'thus the model is extrapolating a result.',
     'function_name': 'high_bmdl',
     'name': 'BMDL higher',
     'threshold': 1,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 0,
     'description': 'The BMD estimate is lower than the lowest non-zero dose '
                    ' (caution).',
     'function_name': 'low_bmd',
     'name': 'Low BMD (warning)',
     'threshold': 3,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 0,
     'description': 'The BMDL estimate is lower than the lowest non-zero dose '
                    ' (caution).',
     'function_name': 'low_bmdl',
     'name': 'Low BMDL (warning)',
     'threshold': 3,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 1,
     'description': 'The BMD estimate is lower than the lowest non-zero dose '
                    ' (warning).',
     'function_name': 'low_bmd',
     'name': 'Low BMD (serious)',
     'threshold': 10,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 1,
     'description': 'The BMDL estimate is lower than the lowest non-zero dose '
                    ' (warning).',
     'function_name': 'low_bmdl',
     'name': 'Low BMDL (serious)',
     'threshold': 10,
     'continuous_on': True,
     'dichotomous_on': True,
     'cancer_dichotomous_on': True},
    {'failure_bin': 1,
     'description': 'The absolute value of the residual at the control is large, '
                    'which is often used to estimate the BMR. Thus, the BMR '
                    'estimate may be inaccurate.',
     'function_name': 'control_residual',
     'name': 'Control residual',
     'threshold': 2,
     'continuous_on': True,
     'dichotomous_on': False,
     'cancer_dichotomous_on': False},
    {'failure_bin': 1,
     'description': 'The standard deviation estimate at the control is '
                    'different the the reported deviation, which is often used '
                    'to estimate the BMR. Thus, the BMR estimate may be inaccurate.',
     'function_name': 'control_stdev',
     'name': 'Control stdev',
     'threshold': 1.5,
     'continuous_on': True,
     'dichotomous_on': False,
     'cancer_dichotomous_on': False}
]
