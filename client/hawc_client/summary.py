import pathlib

import pandas as pd
from docx import Document
from docx.shared import Inches
from playwright.sync_api import Page, sync_playwright

from .client import BaseClient


class SummaryClient(BaseClient):
    """
    Client class for summary requests.
    """

    def visual_list(self, assessment_id: int) -> pd.DataFrame:
        """
        Retrieves a visual list for the given assessment.

        Args:
            assessment_id (int): Assessment ID

        Returns:
            pd.DataFrame: Visual list
        """
        url = f"{self.session.root_url}/summary/api/visual/?assessment_id={assessment_id}"
        response_json = self.session.get(url).json()
        return pd.DataFrame(response_json)

    def datapivot_list(self, assessment_id: int) -> pd.DataFrame:
        """
        Retrieves a data pivot list for the given assessment.

        Args:
            assessment_id (int): Assessment ID

        Returns:
            pd.DataFrame: Data Pivot list
        """
        url = f"{self.session.root_url}/summary/api/data_pivot/?assessment_id={assessment_id}"
        response_json = self.session.get(url).json()
        return pd.DataFrame(response_json)

    def table_list(self, assessment_id: int) -> pd.DataFrame:
        """
        Retrieves a table list for the given assessment.

        Args:
            assessment_id (int): Assessment ID

        Returns:
            pd.DataFrame: Data Pivot list
        """
        url = f"{self.session.root_url}/summary/api/summary-table/?assessment_id={assessment_id}"
        response_json = self.session.get(url).json()
        return pd.DataFrame(response_json)

    def download_visuals(
        self, assessment_id: int, dest: str = ".", file_type: str = "docx", headless: bool = True
    ):
        """
        Retrieves all visuals for the given assessment and saves them in the given path.
        Session must be authenticated with a token to access private assessment visuals.

        Args:
            assessment_id (int): Assessment ID
            dest (str): Folder where images will be saved. Defaults to current working directory.
            file_type (str): 'docx' will save all images to a Word document in the dest folder.
                'png' will save all images separately as pngs in the dest folder. Defaults to 'docx'
            headless (bool): If False, will display the browser while downloading images. Defaults
                to True.

            Returns:
                None
        """

        def _prepare_download(page: Page, row):
            page.goto(row.full_url)

            if row.visual_type == "embedded external website":
                # assume embedded external website is tableau
                download_button = page.frame_locator("iframe").locator(
                    'div[role="button"]:has-text("Download")'
                )
                download_confirm = page.frame_locator("iframe").locator('button:has-text("Image")')
            else:
                download_button = page.locator("button", has=page.locator("i.fa-download"))
                download_confirm = page.locator("text=Download as a PNG")

            download_button.click()

            return download_confirm

        def _save_to_pngs(data: pd.DataFrame, page: Page, dest_path: pathlib.Path):
            for row in data[["slug", "full_url", "visual_type"]].itertuples():
                try:
                    download_confirm = _prepare_download(page, row)
                    # Start waiting for the download
                    with page.expect_download() as download_info:
                        download_confirm.click()
                        download = download_info.value
                        # Save downloaded file
                        download.save_as(str(dest_path.joinpath(f"{row.slug}.png")))
                except Exception as e:
                    # create a txt file with error info
                    with open(dest_path.joinpath(f"{row.slug}-error.txt"), "w") as f:
                        f.write(f"There was an error getting the visual named {row.slug}.\n{e}")

        def _save_to_docx(data: pd.DataFrame, page: Page, dest_path: pathlib.Path):
            doc = Document()
            for row in data[["slug", "full_url", "visual_type"]].itertuples():
                # path for tmp image file
                path = dest_path.joinpath(f"tmp/{row.slug}.png")

                try:
                    download_confirm = _prepare_download(page, row)
                    # Start waiting for the download
                    with page.expect_download() as download_info:
                        download_confirm.click()
                        download = download_info.value
                        # create tmp png file
                        download.save_as(str(path))
                        # add png to doc
                        doc.add_picture(str(path), width=Inches(5.5))
                except Exception as e:
                    # add error info
                    doc.add_paragraph(
                        f"There was an error getting the visual named {row.slug}.\n{e}"
                    )

                # remove tmp file if created
                path.unlink(missing_ok=True)

            doc.save(dest_path.joinpath(f"assessment-{assessment_id}-visuals.docx"))
            # remove tmp folder
            dest_path.joinpath("tmp").rmdir()

        valid_file_types = ["png", "docx"]

        if file_type not in valid_file_types:
            raise ValueError(f"Valid file types are {valid_file_types}")

        dest_path = pathlib.Path(dest)
        visuals = self.visual_list(assessment_id)
        dp = self.datapivot_list(assessment_id)
        data = pd.concat([visuals, dp])
        root_url = self.session.root_url

        if data.shape[0] == 0:
            raise ValueError("No visuals found for given assessment id")

        data.loc[:, "full_url"] = root_url + data.url

        with sync_playwright() as p:
            browser = p.chromium.launch(headless=headless)
            page = browser.new_page()

            # Give playwright authorization to view private pages
            token = self.session._session.headers.get("Authorization")
            if token:
                page.set_extra_http_headers({"Authorization": str(token)})
                page.goto(f"{root_url}/user/api/validate-token/?login=1")

            if file_type == "docx":
                _save_to_docx(data, page, dest_path)
            elif file_type == "png":
                _save_to_pngs(data, page, dest_path)
            browser.close()
